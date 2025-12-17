"""Insert comments and tracked changes into Word documents and Google Docs."""

from typing import List, Dict, Optional
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
import os
import pickle
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime

# Try to import docx, but handle lxml import errors gracefully
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"WARNING: python-docx is not available: {e}")
    print("Word document processing will be limited.")
    # Create dummy classes to prevent errors
    class Document:
        def __init__(self, *args, **kwargs):
            raise ImportError("python-docx is not available. lxml installation failed.")
    class OxmlElement:
        pass
    def qn(*args, **kwargs):
        raise ImportError("python-docx is not available. lxml installation failed.")

# Try to import lxml, but provide a fallback if it's not available
try:
    from lxml import etree as lxml_etree
    LXML_AVAILABLE = True
except ImportError as e:
    LXML_AVAILABLE = False
    print(f"WARNING: lxml is not available: {e}")
    print("Word document formatting may be affected. Install lxml for best results.")
    # Create a dummy lxml_etree object to prevent errors
    class DummyLxmlEtree:
        @staticmethod
        def fromstring(xml_bytes):
            return ET.fromstring(xml_bytes)
        
        @staticmethod
        def tostring(element, encoding='utf-8', pretty_print=False, xml_declaration=False):
            # Fallback to ElementTree's tostring
            if isinstance(element, bytes):
                element = ET.fromstring(element)
            return ET.tostring(element, encoding=encoding)
        
        class XMLSyntaxError(Exception):
            pass
    
    lxml_etree = DummyLxmlEtree()


class CommentInserter:
    """Inserts comments into documents based on redline analysis."""
    
    def __init__(self, doc_path: str = None, doc_id: str = None, credentials_path: str = None):
        """Initialize with document path or Google Doc ID."""
        self.doc_path = doc_path
        self.doc_id = doc_id
        self.credentials_path = credentials_path or os.getenv('GOOGLE_CREDENTIALS_PATH', 'credentials.json')
        self.document = None
        self.service = None
        
        if doc_path:
            self.document = Document(doc_path)
        elif doc_id:
            self.service = self._authenticate_google()
    
    def _authenticate_google(self):
        """Authenticate for Google Docs API."""
        from google_auth_oauthlib.flow import InstalledAppFlow
        from google.auth.transport.requests import Request
        
        SCOPES = ['https://www.googleapis.com/auth/documents']
        creds = None
        
        if os.path.exists('token.pickle'):
            with open('token.pickle', 'rb') as token:
                creds = pickle.load(token)
        
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                if not os.path.exists(self.credentials_path):
                    raise FileNotFoundError(f"Google credentials not found at {self.credentials_path}")
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_path, SCOPES)
                creds = flow.run_local_server(port=0)
            
            with open('token.pickle', 'wb') as token:
                pickle.dump(creds, token)
        
        return build('docs', 'v1', credentials=creds)
    
    def insert_comments_word(self, analyses: List[Dict], output_path: str, use_tracked_changes: bool = False, extractor=None) -> None:
        """Insert comments into Word document and optionally add counter redlines.
        
        Args:
            analyses: List of analysis results
            output_path: Path to save the output document
            use_tracked_changes: When True, add a tracked-change insertion with the recommended action
            extractor: WordRedlineExtractor instance to access XML element references
        """
        if not self.document:
            raise ValueError("Word document not loaded")
        
        import sys
        print("\n" + "="*80, flush=True)
        print("=== INSERTING COMMENTS INTO WORD DOCUMENT ===", flush=True)
        print(f"Number of analyses: {len(analyses)}", flush=True)
        print(f"Output path: {output_path}", flush=True)
        print("="*80 + "\n", flush=True)
        
        # Use native Word comments (comment bubbles in comments pane)
        print("Calling _insert_comments_using_docx_api...", flush=True)
        self._insert_comments_using_docx_api(
            analyses,
            output_path,
            extractor,
            use_tracked_changes=use_tracked_changes
        )
        print("\n=== COMMENT INSERTION COMPLETE ===\n", flush=True)
    
    def _insert_comments_as_annotations(self, analyses: List[Dict], output_path: str, extractor=None) -> None:
        """Insert comments as formatted annotations in the document, positioned after each redline."""
        # If we have extractor with XML element references, use XML method for precise positioning
        if extractor and hasattr(extractor, 'doc_path'):
            self._insert_comments_via_xml(analyses, output_path, extractor)
        else:
            # Fallback to simple paragraph-based method
            commented_paragraphs = set()
            
            for analysis in analyses:
                redline = analysis['redline']
                comment_text = analysis.get('comment_text', analysis.get('response', ''))
                risk_level = analysis.get('risk_level', 'Medium')
                
                if not comment_text:
                    continue
                
                target_text = redline.get('text', '')
                if not target_text:
                    continue
                
                # Search for text in document
                for para_idx, paragraph in enumerate(self.document.paragraphs):
                    if para_idx in commented_paragraphs:
                        continue
                        
                    if target_text[:50] in paragraph.text:
                        self._add_comment_annotation(paragraph, comment_text, risk_level)
                        commented_paragraphs.add(para_idx)
                        break
            
            self.document.save(output_path)
    
    def _insert_comments_via_xml(self, analyses: List[Dict], output_path: str, extractor) -> None:
        """Insert comments via XML manipulation for precise positioning."""
        temp_path = output_path.replace('.docx', '_temp.docx')
        self.document.save(temp_path)
        
        with zipfile.ZipFile(temp_path, 'r') as docx:
            document_xml = docx.read('word/document.xml')
            root = ET.fromstring(document_xml)
            
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Process each analysis and insert comment right after the redline
            for analysis in analyses:
                redline = analysis['redline']
                comment_text = analysis.get('comment_text', analysis.get('response', ''))
                risk_level = analysis.get('risk_level', 'Medium')
                
                if not comment_text:
                    continue
                
                # Find the redline element in the current XML tree by matching text and type
                redline_type = redline.get('type', '')
                redline_text = redline.get('text', '')
                
                if redline_type and redline_text:
                    # Find matching element in current XML
                    redline_elem = self._find_redline_element_in_xml(root, redline_type, redline_text, namespaces)
                    
                    if redline_elem is not None:
                        # Find the parent paragraph by traversing up the tree
                        parent = self._find_parent_paragraph(root, redline_elem, namespaces)
                        
                        if parent is not None:
                            # Insert comment annotation right after the redline element
                            self._insert_comment_after_element(parent, redline_elem, comment_text, risk_level, namespaces)
        
        # Write updated document
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as new_docx:
            with zipfile.ZipFile(temp_path, 'r') as original:
                for item in original.infolist():
                    if item.filename != 'word/document.xml':
                        new_docx.writestr(item, original.read(item.filename))
            
            new_docx.writestr('word/document.xml', ET.tostring(root, encoding='utf-8', xml_declaration=True))
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    def _insert_comments_via_xml_direct(self, analyses: List[Dict], output_path: str, extractor, root: ET.Element, namespaces: Dict, temp_path: str, use_tracked_changes: bool = False) -> None:
        """Insert native Word comments via direct XML manipulation.
        
        This creates proper Word comment bubbles by:
        1. Creating comment entries in word/comments.xml
        2. Adding comment range markers in document.xml
        3. Linking them together with comment IDs
        """
        print(f"\n=== Starting XML-based native comment insertion for {len(analyses)} redlines ===")
        
        # Read comments.xml if it exists, or create it
        # CRITICAL: Register namespace BEFORE creating elements to get 'w:' prefix
        ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
        
        with zipfile.ZipFile(temp_path, 'r') as docx:
            try:
                comments_xml = docx.read('word/comments.xml')
                comments_root = ET.fromstring(comments_xml)
            except KeyError:
                # Create new comments.xml with proper namespace
                comments_root = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments')
        
        # Track comment IDs
        comment_id = 0
        comments_added = 0
        processed_redline_ids = set()
        
        # Process EACH analysis individually - one comment per redline
        for analysis_idx, analysis in enumerate(analyses):
            redline = analysis['redline']
            print(f"Processing redline {analysis_idx + 1} of {len(analyses)}: {redline.get('type', 'unknown')} - {redline.get('text', '')[:50]}...")
            
            # Get all available guidance text
            playbook_principle = analysis.get('playbook_principle', '')
            assessment = analysis.get('assessment', '')
            response = analysis.get('response', '')
            fallbacks = analysis.get('fallbacks', '')
            comment_text = analysis.get('comment_text', '')
            risk_level = analysis.get('risk_level', 'Medium')
            
            # Combine guidance to match summary output format exactly
            # This ensures Word comments match what's shown in the summary
            # CRITICAL: Always reference the playbook principle FIRST
            guidance_parts = []
            
            # Risk Level (shown in summary as badge)
            guidance_parts.append(f"RISK LEVEL: {risk_level}")
            guidance_parts.append("=" * 50)
            
            # PLAYBOOK REFERENCE FIRST - Always cite the playbook principle before any analysis
            if playbook_principle:
                guidance_parts.append(f"\nPLAYBOOK REFERENCE:\n{playbook_principle}")
            
            # Combine assessment and comment_text into a single Assessment field
            combined_assessment = ""
            if assessment and comment_text:
                combined_assessment = f"{assessment}\n\n{comment_text}"
            elif comment_text:
                combined_assessment = comment_text
            else:
                combined_assessment = assessment
            
            # Assessment (combined with comment_text)
            if combined_assessment:
                guidance_parts.append(f"\nASSESSMENT:\n{combined_assessment}")
            
            # Recommended Action (matches summary output - shown as "Recommended Action")
            if response:
                guidance_parts.append(f"\nRECOMMENDED ACTION:\n{response}")
            
            # If no guidance was found, provide default message
            if len(guidance_parts) <= 3:  # Only risk level, separator, and maybe playbook
                guidance_parts.append("\nPlease review this change against the legal playbook.")
            
            full_guidance = "\n".join(guidance_parts)
            
            if not full_guidance.strip():
                print(f"  Skipping - no guidance text")
                continue
            
            # Find the redline element in XML
            redline_type = redline.get('type', '')
            redline_text = redline.get('text', '').strip()
            
            if redline_type not in ['insertion', 'deletion', 'replacement']:
                print(f"  ⚠ REJECTED: Redline type '{redline_type}' is not a tracked change.")
                continue
            
            redline_elem = None
            
            # Handle replacements differently - search for deletion element first (what was replaced)
            if redline_type == 'replacement':
                old_text = redline.get('old_text', '').strip()
                new_text = redline.get('new_text', '').strip()
                
                # Try to find deletion element first (preferred for comment placement)
                if old_text:
                    all_matching_redlines = self._find_all_redline_elements_in_xml(root, 'deletion', old_text, namespaces)
                    print(f"  Found {len(all_matching_redlines)} matching deletion element(s) for replacement")
                    for candidate in all_matching_redlines:
                        elem_id = self._get_element_identifier(candidate, root, namespaces)
                        if elem_id and elem_id not in processed_redline_ids:
                            redline_elem = candidate
                            processed_redline_ids.add(elem_id)
                            print(f"  Selected deletion element for replacement")
                            break
                
                # Fallback: if deletion not found, try insertion element
                if redline_elem is None and new_text:
                    all_matching_redlines = self._find_all_redline_elements_in_xml(root, 'insertion', new_text, namespaces)
                    print(f"  Found {len(all_matching_redlines)} matching insertion element(s) for replacement")
                    for candidate in all_matching_redlines:
                        elem_id = self._get_element_identifier(candidate, root, namespaces)
                        if elem_id and elem_id not in processed_redline_ids:
                            redline_elem = candidate
                            processed_redline_ids.add(elem_id)
                            print(f"  Selected insertion element for replacement")
                            break
            else:
                # For insertion/deletion, use existing search logic
                if not redline_text:
                    print(f"  ⚠ REJECTED: Missing text for {redline_type}")
                    continue
                
                # Find ALL matching redline elements, then pick one we haven't processed
                all_matching_redlines = self._find_all_redline_elements_in_xml(root, redline_type, redline_text, namespaces)
                print(f"  Found {len(all_matching_redlines)} matching redline element(s) in XML")
                
                for candidate in all_matching_redlines:
                    elem_id = self._get_element_identifier(candidate, root, namespaces)
                    if elem_id and elem_id not in processed_redline_ids:
                        redline_elem = candidate
                        processed_redline_ids.add(elem_id)
                        print(f"  Selected redline element: {elem_id}")
                        break
            
            if redline_elem is None:
                print(f"  ⚠ Skipping - could not find unprocessed redline element")
                continue
            
            # Find parent paragraph
            parent_para = self._find_parent_paragraph(root, redline_elem, namespaces)
            if not parent_para:
                print(f"  ⚠ Could not find parent paragraph - skipping")
                continue
            
            # Create comment ID
            comment_id += 1
            
            # Auto-redline based on AI recommendation
            # NOTE: We do NOT duplicate the counterparty's redline - their changes are already 
            # visible in tracked changes. We only insert our counter-proposal text.
            auto_action = analysis.get('auto_redline_action', 'comment_only')
            auto_text = analysis.get('auto_redline_text', '')
            
            if use_tracked_changes and auto_action in ['reject_restore', 'reject_replace']:
                try:
                    # Insert our replacement/restoration text (counterparty's redline is already visible)
                    if auto_action == 'reject_restore':
                        # Restore the original/deleted text
                        if redline_type == 'deletion':
                            # Restore the deleted text
                            restore_text = redline.get('text', '')
                            if restore_text:
                                self._insert_auto_redline(parent_para, redline_elem, restore_text, namespaces, is_restore=True)
                                print(f"  ✓ Auto-redline: Inserted restoration of deleted text")
                        elif redline_type == 'replacement':
                            # Restore the old text that was replaced
                            restore_text = redline.get('old_text', '')
                            if restore_text:
                                self._insert_auto_redline(parent_para, redline_elem, restore_text, namespaces, is_restore=True)
                                print(f"  ✓ Auto-redline: Inserted original text '{restore_text[:50]}...'")
                        elif redline_type == 'insertion':
                            # For unwanted insertions, add a deletion marker
                            ins_text = redline.get('text', '')
                            if ins_text:
                                self._insert_tracked_deletion(parent_para, redline_elem, ins_text, namespaces)
                                print(f"  ✓ Auto-redline: Marked counterparty insertion for deletion")
                    elif auto_action == 'reject_replace':
                        if auto_text:
                            # Insert specific replacement text from playbook
                            self._insert_auto_redline(parent_para, redline_elem, auto_text, namespaces, is_restore=False)
                            print(f"  ✓ Auto-redline: Inserted playbook text '{auto_text[:50]}...'")
                        elif redline_type == 'insertion':
                            # For unwanted insertions without replacement, add a deletion marker
                            ins_text = redline.get('text', '')
                            if ins_text:
                                self._insert_tracked_deletion(parent_para, redline_elem, ins_text, namespaces)
                                print(f"  ✓ Auto-redline: Marked counterparty insertion for deletion")
                except Exception as e:
                    print(f"  ⚠ Could not add auto-redline: {e}")
            elif auto_action == 'accept':
                print(f"  ℹ Auto-redline: Change accepted per playbook (no counter-redline needed)")
            
            # Create the comment in comments.xml
            self._create_word_comment(parent_para, redline_elem, comments_root, comment_id, full_guidance, risk_level, namespaces)
            
            comments_added += 1
            print(f"  ✓ Added comment #{comments_added} to redline")
        
        # Validate before saving
        print(f"\n=== Validating document structure before saving ===")
        validation_errors = self._validate_word_document_structure(root, comments_root, namespaces)
        
        if validation_errors:
            print(f"⚠ WARNING: Found {len(validation_errors)} validation issue(s):", flush=True)
            for error in validation_errors:
                print(f"  - {error}", flush=True)
        else:
            print("✓ Document structure validation passed", flush=True)
        
        # Write updated document
        print(f"\n=== Saving document with {comments_added} comments added ===")
        
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as new_docx:
            # Copy all files from original (we'll handle specific files separately)
            with zipfile.ZipFile(temp_path, 'r') as original:
                for item in original.infolist():
                    if item.filename not in ['word/document.xml', 'word/comments.xml', 'word/_rels/document.xml.rels', '[Content_Types].xml']:
                        new_docx.writestr(item, original.read(item.filename))
            
            # CRITICAL: Ensure document.xml.rels exists and links to comments.xml
            # This relationship file is REQUIRED for Word to recognize comments
            try:
                with zipfile.ZipFile(temp_path, 'r') as original:
                    try:
                        rels_xml = original.read('word/_rels/document.xml.rels')
                        # Parse and check if comments relationship exists
                        rels_root = ET.fromstring(rels_xml)
                        namespaces_rels = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                        
                        # Check if comments relationship already exists
                        has_comments_rel = False
                        for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship', namespaces_rels):
                            rel_type = rel.get('Type', '')
                            if 'comments' in rel_type.lower():
                                has_comments_rel = True
                                break
                        
                        if not has_comments_rel and comments_added > 0:
                            # Add comments relationship
                            print(f"    Adding comments relationship to document.xml.rels", flush=True)
                            # Find the highest relationship ID
                            max_id = 0
                            for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship', namespaces_rels):
                                rel_id = rel.get('Id', '')
                                if rel_id.startswith('rId'):
                                    try:
                                        id_num = int(rel_id[3:])
                                        max_id = max(max_id, id_num)
                                    except ValueError:
                                        pass
                            
                            # Create new relationship
                            new_rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
                            new_rel.set('Id', f'rId{max_id + 1}')
                            new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
                            new_rel.set('Target', 'comments.xml')
                            rels_root.append(new_rel)
                            
                            # Write updated rels file using lxml for proper formatting
                            try:
                                rels_xml_bytes = ET.tostring(rels_root, encoding='utf-8', xml_declaration=True, method='xml')
                                rels_lxml_root = lxml_etree.fromstring(rels_xml_bytes)
                                rels_xml_str = lxml_etree.tostring(
                                    rels_lxml_root,
                                    encoding='utf-8',
                                    xml_declaration=True,
                                    pretty_print=False,
                                    method='xml'
                                ).decode('utf-8')
                                new_docx.writestr('word/_rels/document.xml.rels', rels_xml_str.encode('utf-8'))
                            except Exception as e:
                                # Fallback to ElementTree
                                ET.register_namespace('r', 'http://schemas.openxmlformats.org/package/2006/relationships')
                                rels_xml_str = ET.tostring(rels_root, encoding='utf-8', xml_declaration=True, method='xml')
                                new_docx.writestr('word/_rels/document.xml.rels', rels_xml_str)
                            print(f"    ✓ Added comments relationship", flush=True)
                        else:
                            # Use existing rels file
                            new_docx.writestr('word/_rels/document.xml.rels', rels_xml)
                    except KeyError:
                        # Create new rels file if it doesn't exist
                        if comments_added > 0:
                            print(f"    Creating word/_rels/document.xml.rels with comments relationship", flush=True)
                            rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>
</Relationships>'''
                            new_docx.writestr('word/_rels/document.xml.rels', rels_xml)
                            print(f"    ✓ Created comments relationship file", flush=True)
            except Exception as e:
                print(f"    ⚠ Error handling relationship file: {e}", flush=True)
                # Try to create it anyway if comments were added
                if comments_added > 0:
                    try:
                        rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>
</Relationships>'''
                        new_docx.writestr('word/_rels/document.xml.rels', rels_xml)
                        print(f"    ✓ Created comments relationship file (fallback)", flush=True)
                    except:
                        pass
            
            # CRITICAL: Ensure [Content_Types].xml has an override for comments.xml
            # This tells Word that comments.xml is a valid document part
            # Word REQUIRES this to recognize and display comments
            try:
                with zipfile.ZipFile(temp_path, 'r') as original:
                    try:
                        content_types_xml = original.read('[Content_Types].xml')
                        content_types_root = ET.fromstring(content_types_xml)
                        namespaces_ct = {'ct': 'http://schemas.openxmlformats.org/package/2006/content-types'}
                        
                        # Check if comments.xml override already exists
                        has_comments_override = False
                        for override in content_types_root.findall('.//{http://schemas.openxmlformats.org/package/2006/content-types}Override'):
                            part_name = override.get('PartName', '')
                            if part_name == '/word/comments.xml':
                                has_comments_override = True
                                break
                        
                        if not has_comments_override and comments_added > 0:
                            print(f"    Adding comments.xml override to [Content_Types].xml", flush=True)
                            # Create new override element
                            new_override = ET.Element('{http://schemas.openxmlformats.org/package/2006/content-types}Override')
                            new_override.set('PartName', '/word/comments.xml')
                            new_override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
                            content_types_root.append(new_override)
                            
                            # Write updated Content_Types.xml using lxml for proper formatting
                            try:
                                ct_xml_bytes = ET.tostring(content_types_root, encoding='utf-8', xml_declaration=True, method='xml')
                                ct_lxml_root = lxml_etree.fromstring(ct_xml_bytes)
                                ct_xml_str = lxml_etree.tostring(
                                    ct_lxml_root,
                                    encoding='utf-8',
                                    xml_declaration=True,
                                    pretty_print=False,
                                    method='xml'
                                ).decode('utf-8')
                                new_docx.writestr('[Content_Types].xml', ct_xml_str.encode('utf-8'))
                                print(f"    ✓ Added comments.xml override to [Content_Types].xml", flush=True)
                            except Exception as e:
                                # Fallback to ElementTree
                                ET.register_namespace('ct', 'http://schemas.openxmlformats.org/package/2006/content-types')
                                ct_xml_str = ET.tostring(content_types_root, encoding='utf-8', xml_declaration=True, method='xml')
                                new_docx.writestr('[Content_Types].xml', ct_xml_str)
                                print(f"    ✓ Added comments.xml override (using ElementTree fallback)", flush=True)
                        else:
                            # Use existing Content_Types.xml
                            new_docx.writestr('[Content_Types].xml', content_types_xml)
                            if has_comments_override:
                                print(f"    ✓ [Content_Types].xml already has comments.xml override", flush=True)
                    except KeyError:
                        # Create new Content_Types.xml if it doesn't exist (shouldn't happen, but handle it)
                        if comments_added > 0:
                            print(f"    Creating [Content_Types].xml with comments.xml override", flush=True)
                            content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>'''
                            new_docx.writestr('[Content_Types].xml', content_types_xml.encode('utf-8'))
                            print(f"    ✓ Created [Content_Types].xml with comments override", flush=True)
            except Exception as e:
                print(f"    ⚠ Error handling [Content_Types].xml: {e}", flush=True)
                import traceback
                print(f"    Traceback: {traceback.format_exc()}", flush=True)
                # Fallback: try to create it if comments were added
                if comments_added > 0:
                    try:
                        with zipfile.ZipFile(temp_path, 'r') as original:
                            # Try to read existing one first
                            try:
                                content_types_xml = original.read('[Content_Types].xml')
                                # Parse and add override
                                content_types_root = ET.fromstring(content_types_xml)
                                new_override = ET.Element('{http://schemas.openxmlformats.org/package/2006/content-types}Override')
                                new_override.set('PartName', '/word/comments.xml')
                                new_override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
                                content_types_root.append(new_override)
                                ct_xml_str = ET.tostring(content_types_root, encoding='utf-8', xml_declaration=True, method='xml')
                                new_docx.writestr('[Content_Types].xml', ct_xml_str)
                                print(f"    ✓ Added comments.xml override (fallback method)", flush=True)
                            except KeyError:
                                # Create from scratch
                                content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>'''
                                new_docx.writestr('[Content_Types].xml', content_types_xml.encode('utf-8'))
                                print(f"    ✓ Created [Content_Types].xml with comments override (fallback)", flush=True)
                    except Exception as e2:
                        print(f"    ✗ ERROR: Could not create [Content_Types].xml: {e2}", flush=True)
            
            # Write updated document.xml with proper formatting using lxml for better XML handling
            # CRITICAL: Use lxml consistently for all XML operations to ensure Word compatibility
            try:
                # Convert ElementTree to lxml for proper namespace handling
                doc_xml_bytes = ET.tostring(root, encoding='utf-8', xml_declaration=True, method='xml')
                # Parse with lxml to ensure valid XML and proper structure
                doc_lxml_root = lxml_etree.fromstring(doc_xml_bytes)
                
                # CRITICAL: Ensure all text content is properly escaped
                # lxml handles this automatically, but we'll validate
                for elem in doc_lxml_root.iter():
                    if elem.text:
                        # lxml will automatically escape special characters
                        pass
                
                # Write with lxml to ensure proper formatting and namespace prefixes
                # CRITICAL: Use method='xml' and ensure proper encoding
                doc_xml_str = lxml_etree.tostring(
                    doc_lxml_root, 
                    encoding='utf-8', 
                    xml_declaration=True, 
                    pretty_print=False,
                    method='xml'
                ).decode('utf-8')
                
                # Ensure 'w:' prefix is used (lxml should handle this, but double-check)
                if 'ns0:' in doc_xml_str:
                    print(f"    ⚠ Fixing namespace prefix in document.xml: replacing ns0: with w:", flush=True)
                    doc_xml_str = doc_xml_str.replace('ns0:', 'w:').replace('xmlns:ns0=', 'xmlns:w=')
                
                # CRITICAL: Validate XML structure before writing
                try:
                    lxml_etree.fromstring(doc_xml_str.encode('utf-8'))
                    print(f"    ✓ document.xml structure validated", flush=True)
                except lxml_etree.XMLSyntaxError as e:
                    print(f"    ⚠ WARNING: document.xml validation error: {e}", flush=True)
                    print(f"    Line {e.lineno}, column {e.offset}", flush=True)
                
                # Write with explicit UTF-8 encoding
                new_docx.writestr('word/document.xml', doc_xml_str.encode('utf-8'))
                print(f"    ✓ Wrote document.xml", flush=True)
            except Exception as e:
                print(f"    ⚠ Error processing document.xml with lxml: {e}", flush=True)
                import traceback
                print(f"    Traceback: {traceback.format_exc()}", flush=True)
                # Fallback: try ElementTree but still validate
                ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                doc_xml_str = ET.tostring(root, encoding='utf-8', xml_declaration=True, method='xml')
                # Validate even the fallback
                try:
                    lxml_etree.fromstring(doc_xml_str)
                    print(f"    ✓ Fallback document.xml validated", flush=True)
                except lxml_etree.XMLSyntaxError as e:
                    print(f"    ✗ ERROR: Fallback document.xml has XML errors: {e}", flush=True)
                new_docx.writestr('word/document.xml', doc_xml_str)
            
            # Write updated comments.xml with proper formatting using lxml
            # Ensure comments root has proper structure
            if len(list(comments_root)) == 0:
                print("⚠ WARNING: No comments to write to comments.xml", flush=True)
            else:
                # CRITICAL: Verify each comment has text before writing
                for comment in comments_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment', namespaces):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    text_elems = comment.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', namespaces)
                    has_text = any(te.text and te.text.strip() for te in text_elems)
                    if not has_text:
                        print(f"    ⚠ WARNING: Comment {comment_id} has no text before writing!", flush=True)
                
                # Use lxml for proper XML generation with correct namespace prefixes
                try:
                    comments_xml_bytes = ET.tostring(comments_root, encoding='utf-8', xml_declaration=True, method='xml')
                    # Parse with lxml to ensure valid XML structure
                    comments_lxml_root = lxml_etree.fromstring(comments_xml_bytes)
                    
                    # CRITICAL: Ensure all text content is properly escaped
                    # lxml handles this automatically, but we'll validate
                    for elem in comments_lxml_root.iter():
                        if elem.text:
                            # lxml will automatically escape special characters like &, <, >
                            # But we need to ensure text is not None
                            if elem.text is None:
                                elem.text = ''
                    
                    # Write with lxml - this ensures proper namespace handling and 'w:' prefix
                    # CRITICAL: Use method='xml' and ensure proper encoding
                    comments_xml_str = lxml_etree.tostring(
                        comments_lxml_root, 
                        encoding='utf-8', 
                        xml_declaration=True, 
                        pretty_print=False,
                        method='xml'
                    ).decode('utf-8')
                    
                    # Ensure 'w:' prefix is used (lxml should handle this, but double-check)
                    if 'ns0:' in comments_xml_str:
                        print(f"    ⚠ Fixing namespace prefix in comments.xml: replacing ns0: with w:", flush=True)
                        comments_xml_str = comments_xml_str.replace('ns0:', 'w:').replace('xmlns:ns0=', 'xmlns:w=')
                    
                    # CRITICAL: Ensure XML declaration is correct format for Word
                    # Word is very picky about the XML declaration format
                    if not comments_xml_str.startswith('<?xml version="1.0" encoding="UTF-8"?>'):
                        # Fix the declaration if needed
                        if comments_xml_str.startswith('<?xml'):
                            # Replace with standard format
                            lines = comments_xml_str.split('\n', 1)
                            if len(lines) > 1:
                                comments_xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + lines[1]
                            else:
                                comments_xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + comments_xml_str
                    
                    # CRITICAL: Validate XML structure before writing
                    try:
                        lxml_etree.fromstring(comments_xml_str.encode('utf-8'))
                        print(f"    ✓ comments.xml structure validated", flush=True)
                    except lxml_etree.XMLSyntaxError as e:
                        print(f"    ⚠ WARNING: comments.xml validation error: {e}", flush=True)
                        print(f"    Line {e.lineno}, column {e.offset}", flush=True)
                    
                    # Write with explicit UTF-8 encoding
                    new_docx.writestr('word/comments.xml', comments_xml_str.encode('utf-8'))
                    print(f"✓ Wrote {len(list(comments_root))} comment(s) to comments.xml", flush=True)
                except Exception as e:
                    print(f"    ⚠ Error processing comments.xml with lxml: {e}", flush=True)
                    import traceback
                    print(f"    Traceback: {traceback.format_exc()}", flush=True)
                    # Fallback to ElementTree but still validate
                    ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                    comments_xml_str = ET.tostring(comments_root, encoding='utf-8', xml_declaration=True, method='xml')
                    # Validate even the fallback
                    try:
                        lxml_etree.fromstring(comments_xml_str)
                        print(f"    ✓ Fallback comments.xml validated", flush=True)
                    except lxml_etree.XMLSyntaxError as e:
                        print(f"    ✗ ERROR: Fallback comments.xml has XML errors: {e}", flush=True)
                    if b'ns0:' in comments_xml_str:
                        comments_xml_str = comments_xml_str.replace(b'ns0:', b'w:').replace(b'xmlns:ns0=', b'xmlns:w=')
                    new_docx.writestr('word/comments.xml', comments_xml_str)
                    print(f"✓ Wrote {len(list(comments_root))} comment(s) to comments.xml (using ElementTree fallback)", flush=True)
        
        print(f"✓ Document saved to: {output_path}")
        
        # Final validation - try to read the document back
        try:
            test_doc = Document(output_path)
            print("✓ Document validation: Successfully opened saved document", flush=True)
        except Exception as e:
            print(f"⚠ WARNING: Could not validate saved document: {e}", flush=True)
        
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    def _validate_word_document_structure(self, root: ET.Element, comments_root: ET.Element, namespaces: Dict) -> List[str]:
        """Validate Word document structure before saving.
        
        Returns list of validation errors (empty if valid).
        """
        errors = []
        
        # Check that all comment IDs in document.xml have corresponding entries in comments.xml
        comment_ids_in_doc = set()
        comment_ids_in_comments = set()
        
        # Find all comment IDs referenced in document.xml
        for comment_range_start in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart', namespaces):
            comment_id = comment_range_start.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if comment_id:
                comment_ids_in_doc.add(comment_id)
        
        # Find all comment IDs in comments.xml
        for comment in comments_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment', namespaces):
            comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if comment_id:
                comment_ids_in_comments.add(comment_id)
        
        # Check for orphaned comment references
        orphaned = comment_ids_in_doc - comment_ids_in_comments
        if orphaned:
            errors.append(f"Found {len(orphaned)} comment reference(s) in document.xml without corresponding entries in comments.xml: {orphaned}")
        
        # Check for comments without references
        unreferenced = comment_ids_in_comments - comment_ids_in_doc
        if unreferenced:
            errors.append(f"Found {len(unreferenced)} comment(s) in comments.xml without references in document.xml: {unreferenced}")
        
        # Validate comment structure - each comment must have at least one paragraph with text
        for comment in comments_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment', namespaces):
            comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            paragraphs = comment.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', namespaces)
            
            if not paragraphs:
                errors.append(f"Comment {comment_id} has no paragraphs")
            else:
                # Check that at least one paragraph has text
                has_text = False
                for para in paragraphs:
                    text_elems = para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', namespaces)
                    for text_elem in text_elems:
                        if text_elem.text and text_elem.text.strip():
                            has_text = True
                            break
                    if has_text:
                        break
                
                if not has_text:
                    errors.append(f"Comment {comment_id} has no text content")
        
        # Validate comment range markers are properly paired
        for comment_id in comment_ids_in_doc:
            starts = root.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}commentRangeStart[@{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}id="{comment_id}"]', namespaces)
            ends = root.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}commentRangeEnd[@{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}id="{comment_id}"]', namespaces)
            refs = root.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}commentReference[@{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}id="{comment_id}"]', namespaces)
            
            if len(starts) != len(ends):
                errors.append(f"Comment {comment_id}: Mismatched commentRangeStart ({len(starts)}) and commentRangeEnd ({len(ends)})")
            
            if len(refs) == 0:
                errors.append(f"Comment {comment_id}: No commentReference found")
        
        return errors
    
    def _get_font_size_from_element(self, element: ET.Element, paragraph_elem: ET.Element, namespaces: Dict) -> str:
        """Extract font size from an element or its parent paragraph. Returns size in half-points (e.g., '24' for 12pt)."""
        # First, try to get size from the element's run properties
        rpr = element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', namespaces)
        if rpr is not None:
            sz = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', namespaces)
            if sz is not None and sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'):
                return sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        
        # Try to get from any run in the paragraph
        for run in paragraph_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r', namespaces):
            rpr = run.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', namespaces)
            if rpr is not None:
                sz = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', namespaces)
                if sz is not None and sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'):
                    return sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        
        # Try paragraph properties
        ppr = paragraph_elem.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr', namespaces)
        if ppr is not None:
            rpr = ppr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', namespaces)
            if rpr is not None:
                sz = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', namespaces)
                if sz is not None and sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'):
                    return sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        
        # Default to 24 (12pt) if nothing found
        return '24'
    
    def _insert_comment_after_element(self, paragraph_elem: ET.Element, target_elem: ET.Element, comment_text: str, risk_level: str, namespaces: Dict) -> None:
        """Insert a comment annotation right after a specific element in a paragraph."""
        # Create a run for the comment
        run_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        
        # Add formatting
        rpr = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        
        # Risk-based color - Professional color scheme
        color_vals = {
            'High': 'C70039',  # Dark Red/Burgundy
            'Medium': 'FF8C00',  # Dark Orange/Amber
            'Low': '006400'  # Dark Green
        }
        color_val = color_vals.get(risk_level, '003366')  # Default Navy Blue
        
        color = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', color_val)
        italic = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i')
        
        # Get font size from surrounding content
        font_size = self._get_font_size_from_element(target_elem, paragraph_elem, namespaces)
        size = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
        size.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', font_size)
        
        rpr.append(italic)
        rpr.append(color)
        rpr.append(size)
        run_elem.append(rpr)
        
        # Add text
        text_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        text_elem.text = f" [AI Analysis - Risk: {risk_level}] {comment_text}"
        run_elem.append(text_elem)
        
        # Find position of target element and insert after it
        children = list(paragraph_elem)
        try:
            target_idx = children.index(target_elem)
            paragraph_elem.insert(target_idx + 1, run_elem)
        except (ValueError, AttributeError):
            # If we can't find the exact element, append to paragraph
            paragraph_elem.append(run_elem)
    
    def _insert_tracked_changes_word(self, analyses: List[Dict], output_path: str, extractor=None) -> None:
        """Insert responses as Word comments using python-docx's built-in comment support.
        
        This tries to use python-docx's comment API first (more reliable), 
        falls back to formatted text annotations if that's not available.
        """
        # Try using python-docx's built-in comment support first
        try:
            # Check for either doc.add_comment() or doc.comments.add_comment()
            has_comment_api = (hasattr(self.document, 'add_comment') or 
                             (hasattr(self.document, 'comments') and hasattr(self.document.comments, 'add_comment')))
            if has_comment_api:
                self._insert_comments_using_docx_api(analyses, output_path, extractor)
                return
        except Exception as e:
            print(f"Note: python-docx comment API not working, using formatted text: {e}")
        
        # Fallback to formatted text annotations
        self._insert_formatted_annotations_fallback(analyses, output_path, extractor)
    
    def _insert_comments_using_docx_api(self, analyses: List[Dict], output_path: str, extractor=None, use_tracked_changes: bool = False) -> None:
        """Insert comments using python-docx's built-in comment API, associating with each redline."""
        print(f"\n=== Starting comment insertion for {len(analyses)} redlines ===")
        
        # We need to work with XML to find the actual redline elements, then find corresponding runs
        # Save and reload to work with XML structure
        temp_path = output_path.replace('.docx', '_temp.docx')
        self.document.save(temp_path)
        
        # Reload to ensure we have fresh document structure
        # Try bayoo-docx first (better comment support), fallback to python-docx
        try:
            import docx as bayoo_docx
            if hasattr(bayoo_docx, 'Document'):
                doc = bayoo_docx.Document(temp_path)
                print("✓ Using bayoo-docx (enhanced comment support)")
            else:
                raise ImportError("bayoo-docx Document not found")
        except (ImportError, AttributeError) as e:
            print(f"⚠ bayoo-docx not available ({e}), using python-docx")
            from docx import Document
            doc = Document(temp_path)
        
        # Read XML to find redline positions
        with zipfile.ZipFile(temp_path, 'r') as docx:
            document_xml = docx.read('word/document.xml')
            root = ET.fromstring(document_xml)
        
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # python-docx doesn't have add_comment - we need to use XML-based insertion
        # This creates native Word comments properly
        print("Using XML-based native Word comment insertion (python-docx doesn't support add_comment)")
        self._insert_comments_via_xml_direct(
            analyses,
            output_path,
            extractor,
            root,
            namespaces,
            temp_path,
            use_tracked_changes=use_tracked_changes
        )
        return
        
        # Track which redline elements we've already processed to avoid duplicates
        # Use a more reliable identifier: the element's position in the XML tree
        processed_redline_ids = set()
        comments_added = 0
        
        # Process EACH analysis individually - one comment per redline
        for analysis_idx, analysis in enumerate(analyses):
            redline = analysis['redline']
            print(f"Processing redline {analysis_idx + 1} of {len(analyses)}: {redline.get('type', 'unknown')} - {redline.get('text', '')[:50]}...")
            
            # Get all available guidance text
            playbook_principle = analysis.get('playbook_principle', '')
            assessment = analysis.get('assessment', '')
            response = analysis.get('response', '')
            fallbacks = analysis.get('fallbacks', '')
            comment_text = analysis.get('comment_text', '')
            risk_level = analysis.get('risk_level', 'Medium')
            
            # Combine guidance to match summary output format exactly
            guidance_parts = []
            
            # Risk Level (shown in summary as badge)
            guidance_parts.append(f"RISK LEVEL: {risk_level}")
            guidance_parts.append("=" * 50)
            
            # Combine assessment and comment_text into a single Assessment field
            combined_assessment = ""
            if assessment and comment_text:
                combined_assessment = f"{assessment}\n\n{comment_text}"
            elif comment_text:
                combined_assessment = comment_text
            else:
                combined_assessment = assessment
            
            # Assessment (combined with comment_text)
            if combined_assessment:
                guidance_parts.append(f"\nASSESSMENT:\n{combined_assessment}")
            
            # Recommended Action (matches summary output - shown as "Recommended Action")
            if response:
                guidance_parts.append(f"\nRECOMMENDED ACTION:\n{response}")
            
            # If no guidance was found, provide default message
            if len(guidance_parts) <= 2:  # Only risk level and separator
                guidance_parts.append("\nPlease review this change against the legal playbook.")
            
            full_guidance = "\n".join(guidance_parts)
            
            if not full_guidance.strip():
                print(f"  Skipping - no guidance text")
                continue
            
            # CRITICAL VALIDATION: Only process actual tracked changes
            # Reject any redline that is not 'insertion' or 'deletion'
            redline_type = redline.get('type', '')
            redline_text = redline.get('text', '').strip()
            
            # Validate that this is an actual tracked change
            if redline_type not in ['insertion', 'deletion']:
                print(f"  ⚠ REJECTED: Redline type '{redline_type}' is not a tracked change. Only 'insertion' and 'deletion' are processed.")
                continue
            
            if not redline_type or not redline_text:
                print(f"  ⚠ REJECTED: Missing type or text")
                continue
            
            # Find ALL matching redline elements, then pick one we haven't processed
            all_matching_redlines = self._find_all_redline_elements_in_xml(root, redline_type, redline_text, namespaces)
            print(f"  Found {len(all_matching_redlines)} matching redline element(s) in XML")
            
            redline_elem = None
            for candidate in all_matching_redlines:
                # Create a unique identifier based on element position and text
                elem_id = self._get_element_identifier(candidate, root, namespaces)
                if elem_id and elem_id not in processed_redline_ids:
                    redline_elem = candidate
                    processed_redline_ids.add(elem_id)
                    print(f"  Selected redline element: {elem_id}")
                    break
            
            # CRITICAL: Only proceed if we found an actual redline element (w:ins or w:del)
            # Do NOT use fallback text matching - that would match regular document text!
            if redline_elem is None:
                print(f"  ⚠ Skipping - could not find unprocessed redline element in XML for: {redline_text[:50]}")
                if len(all_matching_redlines) > 0:
                    print(f"     Note: Found {len(all_matching_redlines)} matching elements but all were already processed")
                continue
            
            # Handle deletions differently - they don't have regular runs
            if redline_type == 'deletion':
                # For deletions, find the parent paragraph containing the deletion
                parent_para = self._find_parent_paragraph(root, redline_elem, namespaces)
                
                if parent_para is not None:
                    # Find a run that comes AFTER the deletion to attach the comment to
                    # This ensures we don't select the whole paragraph
                    target_run = self._find_run_after_deletion(parent_para, redline_elem, doc, namespaces)
                    
                    if target_run:
                        try:
                            comment = doc.add_comment(
                                runs=[target_run],
                                text=full_guidance,
                                author='RedLine Agent',
                                initials='RA'
                            )
                            comments_added += 1
                            print(f"  ✓ Added comment #{comments_added} to deletion (attached to specific run after deletion)")
                        except Exception as e:
                            print(f"  ⚠ Could not add comment to deletion: {e}")
                            print(f"     Skipping to avoid commenting on non-redline text")
                    else:
                        print(f"  ⚠ Could not find specific run after deletion - skipping to avoid commenting on non-redline text")
                else:
                    # Couldn't find parent paragraph - skip rather than fallback
                    print(f"  ⚠ Could not find parent paragraph for deletion - skipping to avoid commenting on non-redline text")
            
            else:
                # Handle insertions - find ONLY the runs that are WITHIN the redline element
                # Get all runs that are direct or indirect children of the <w:ins> element
                runs_in_redline = redline_elem.findall('.//w:r', namespaces)
                
                if runs_in_redline:
                    print(f"  Found {len(runs_in_redline)} run(s) within redline element")
                    # Get the exact text from this redline element
                    redline_full_text = self._get_text_from_element(redline_elem, namespaces)
                    print(f"  Redline text: {redline_full_text[:100]}...")
                    
                    # Find the corresponding python-docx runs by matching ONLY runs that are part of this specific redline
                    # Pass the root XML element so we can find the parent paragraph
                    target_runs = self._find_runs_for_redline_insertion(redline_elem, redline_full_text, doc, root, namespaces)
                    print(f"  Found {len(target_runs)} matching python-docx run(s)")
                    
                    # If we found runs, add comment ONLY to those specific runs
                    if target_runs:
                        try:
                            # Only attach comment to the specific runs that are part of the redline
                            print(f"  Attempting to add comment to {len(target_runs)} run(s)...")
                            comment = doc.add_comment(
                                runs=target_runs,
                                text=full_guidance,
                                author='RedLine Agent',
                                initials='RA'
                            )
                            comments_added += 1
                            print(f"  ✓ SUCCESS: Added comment #{comments_added} to {len(target_runs)} specific run(s) within redline")
                        except Exception as e:
                            print(f"  ✗ ERROR: Could not add comment to runs: {type(e).__name__}: {e}")
                            import traceback
                            print(f"  Traceback: {traceback.format_exc()}")
                    else:
                        # No runs found - this shouldn't happen for valid redlines, but skip rather than fallback
                        print(f"  ⚠ Could not find specific runs for redline - skipping")
                        print(f"     Redline text was: {redline_full_text[:50]}...")
                else:
                    # No runs found in redline element - skip rather than fallback to avoid commenting on regular text
                    print(f"  ⚠ No runs found in redline element - skipping")
        
        # Save document
        print(f"\n=== Saving document with {comments_added} comments added ===")
        doc.save(output_path)
        print(f"✓ Document saved to: {output_path}")
        
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    def _insert_formatted_annotations_fallback(self, analyses: List[Dict], output_path: str, extractor=None) -> None:
        """Fallback: Insert formatted text annotations when comment API not available."""
        # Save document first to ensure we have a file to work with
        temp_path = output_path.replace('.docx', '_temp.docx')
        self.document.save(temp_path)
        
        # Open the docx file (it's a zip archive)
        with zipfile.ZipFile(temp_path, 'r') as docx:
            # Read document.xml
            document_xml = docx.read('word/document.xml')
            root = ET.fromstring(document_xml)
            
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Process each analysis and insert formatted text annotations
            for analysis in analyses:
                redline = analysis['redline']
                # Get all available guidance text
                playbook_principle = analysis.get('playbook_principle', '')
                assessment = analysis.get('assessment', '')
                response = analysis.get('response', '')
                fallbacks = analysis.get('fallbacks', '')
                comment_text = analysis.get('comment_text', '')
                risk_level = analysis.get('risk_level', 'Medium')
                
                # Combine guidance to match summary output format exactly
                guidance_parts = []
                
                # Risk Level
                guidance_parts.append(f"RISK LEVEL: {risk_level}")
                
                # Combine assessment and comment_text into a single Assessment field
                combined_assessment = ""
                if assessment and comment_text:
                    combined_assessment = f"{assessment} {comment_text}"
                elif comment_text:
                    combined_assessment = comment_text
                else:
                    combined_assessment = assessment
                
                # Assessment (combined with comment_text)
                if combined_assessment:
                    guidance_parts.append(f"ASSESSMENT: {combined_assessment}")
                
                # Recommended Action (matches summary output)
                if response:
                    guidance_parts.append(f"RECOMMENDED ACTION: {response}")
                
                if not guidance_parts:
                    guidance_parts.append("Please review this change against the legal playbook.")
                
                full_guidance = " | ".join(guidance_parts)
                
                if not full_guidance.strip():
                    continue
                
                # Find the redline element in the current XML tree by matching text and type
                redline_type = redline.get('type', '')
                redline_text = redline.get('text', '')
                
                if redline_type and redline_text:
                    # Find matching element in current XML
                    redline_elem = self._find_redline_element_in_xml(root, redline_type, redline_text, namespaces)
                    
                    if redline_elem is not None:
                        # Find the parent paragraph
                        parent = self._find_parent_paragraph(root, redline_elem, namespaces)
                        
                        if parent is not None:
                            # Insert formatted text annotation right after the redline
                            self._insert_formatted_annotation(
                                parent,
                                redline_elem,
                                full_guidance,
                                risk_level,
                                namespaces
                            )
        
        # Write updated document
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as new_docx:
            # Copy all files from original
            with zipfile.ZipFile(temp_path, 'r') as original:
                for item in original.infolist():
                    if item.filename != 'word/document.xml':
                        new_docx.writestr(item, original.read(item.filename))
            
            # Write updated document.xml
            doc_xml_str = ET.tostring(root, encoding='utf-8', xml_declaration=True)
            new_docx.writestr('word/document.xml', doc_xml_str)
        
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    def _find_paragraph_with_text(self, root: ET.Element, search_text: str, namespaces: Dict) -> Optional[ET.Element]:
        """Find a paragraph element containing the search text."""
        for para in root.findall('.//w:p', namespaces):
            text_elements = para.findall('.//w:t', namespaces)
            full_text = ''.join([elem.text or '' for elem in text_elements])
            if search_text in full_text:
                return para
        return None
    
    def _find_redline_element_in_xml(self, root: ET.Element, redline_type: str, redline_text: str, namespaces: Dict) -> Optional[ET.Element]:
        """Find a redline element in the XML tree by matching type and text."""
        all_matches = self._find_all_redline_elements_in_xml(root, redline_type, redline_text, namespaces)
        return all_matches[0] if all_matches else None
    
    def _find_all_redline_elements_in_xml(self, root: ET.Element, redline_type: str, redline_text: str, namespaces: Dict) -> List[ET.Element]:
        """Find ALL actual redline XML elements matching the type and text.
        
        CRITICAL: This method ONLY searches for actual tracked change elements:
        - <w:ins> for insertions
        - <w:del> for deletions
        
        It does NOT match regular document text or any non-redline elements.
        """
        matches = []
        
        # CRITICAL: Only search for actual tracked change XML elements
        # Reject any type that is not a tracked change
        if redline_type not in ['insertion', 'deletion']:
            # Note: 'replacement' is handled separately in the calling code
            print(f"  ⚠ Invalid redline type '{redline_type}' - only 'insertion' and 'deletion' are valid for search")
            return matches
        
        # Normalize the search text for better matching
        search_text_normalized = ' '.join(redline_text.split()).strip() if redline_text else ''
        search_text_short = search_text_normalized[:100] if len(search_text_normalized) > 100 else search_text_normalized
        
        # Search for insertion elements - ONLY <w:ins> elements
        if redline_type == 'insertion':
            all_ins = root.findall('.//w:ins', namespaces)
            print(f"    Searching through {len(all_ins)} insertion element(s) in XML")
            for ins_elem in all_ins:
                # Get text ONLY from within the <w:ins> element
                text_elements = ins_elem.findall('.//w:t', namespaces)
                text_parts = []
                for elem in text_elements:
                    if elem.text:
                        text_parts.append(elem.text)
                elem_text = ''.join(text_parts)
                elem_text_normalized = ' '.join(elem_text.split()).strip() if elem_text else ''
                
                # More flexible matching - try multiple strategies
                matched = False
                if search_text_normalized and elem_text_normalized:
                    # Exact match
                    if search_text_normalized == elem_text_normalized:
                        matched = True
                    # Substring match (either direction)
                    elif search_text_normalized in elem_text_normalized or elem_text_normalized in search_text_normalized:
                        matched = True
                    # First 50 chars match
                    elif (len(search_text_normalized) >= 10 and len(elem_text_normalized) >= 10 and
                          (search_text_normalized[:50] in elem_text_normalized or elem_text_normalized[:50] in search_text_normalized)):
                        matched = True
                    # First word match (for very short redlines)
                    elif len(search_text_normalized) < 10:
                        first_word = search_text_normalized.split()[0] if search_text_normalized.split() else ''
                        if first_word and first_word in elem_text_normalized:
                            matched = True
                
                if matched:
                    matches.append(ins_elem)
                    print(f"      ✓ Matched insertion: '{elem_text_normalized[:50]}...'")
        
        # Search for deletion elements - ONLY <w:del> elements
        elif redline_type == 'deletion':
            all_del = root.findall('.//w:del', namespaces)
            print(f"    Searching through {len(all_del)} deletion element(s) in XML")
            for del_elem in all_del:
                # Get text ONLY from within the <w:del> element
                text_elements = del_elem.findall('.//w:delText', namespaces)
                text_parts = []
                for elem in text_elements:
                    if elem.text:
                        text_parts.append(elem.text)
                elem_text = ''.join(text_parts)
                elem_text_normalized = ' '.join(elem_text.split()).strip() if elem_text else ''
                
                # More flexible matching - try multiple strategies
                matched = False
                if search_text_normalized and elem_text_normalized:
                    # Exact match
                    if search_text_normalized == elem_text_normalized:
                        matched = True
                    # Substring match (either direction)
                    elif search_text_normalized in elem_text_normalized or elem_text_normalized in search_text_normalized:
                        matched = True
                    # First 50 chars match
                    elif (len(search_text_normalized) >= 10 and len(elem_text_normalized) >= 10 and
                          (search_text_normalized[:50] in elem_text_normalized or elem_text_normalized[:50] in search_text_normalized)):
                        matched = True
                    # First word match (for very short redlines)
                    elif len(search_text_normalized) < 10:
                        first_word = search_text_normalized.split()[0] if search_text_normalized.split() else ''
                        if first_word and first_word in elem_text_normalized:
                            matched = True
                
                if matched:
                    matches.append(del_elem)
                    print(f"      ✓ Matched deletion: '{elem_text_normalized[:50]}...'")
        
        print(f"    Found {len(matches)} matching element(s) for redline: '{search_text_short}...'")
        return matches
    
    def _get_element_identifier(self, element: ET.Element, root: ET.Element, namespaces: Dict) -> Optional[str]:
        """Create a unique identifier for an element based on its position and content."""
        # Get the element's tag and text to create a unique identifier
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        text = self._get_text_from_element(element, namespaces)
        
        # Find the element's position by counting similar elements before it
        if tag == 'ins':
            all_ins = root.findall('.//w:ins', namespaces)
            position = list(all_ins).index(element) if element in all_ins else -1
            return f"ins_{position}_{text[:30]}"
        elif tag == 'del':
            all_del = root.findall('.//w:del', namespaces)
            position = list(all_del).index(element) if element in all_del else -1
            return f"del_{position}_{text[:30]}"
        
        return None
    
    def _get_text_from_element(self, element: ET.Element, namespaces: Dict) -> str:
        """Get all text content from an XML element, normalizing whitespace."""
        text_parts = []
        
        # For insertions, get text from w:t elements
        for text_elem in element.findall('.//w:t', namespaces):
            if text_elem.text:
                text_parts.append(text_elem.text)
        
        # For deletions, get text from w:delText elements
        for del_text_elem in element.findall('.//w:delText', namespaces):
            if del_text_elem.text:
                text_parts.append(del_text_elem.text)
        
        text = ''.join(text_parts)
        # Normalize whitespace to match extraction logic
        text = ' '.join(text.split()) if text.strip() else text
        return text.strip()
    
    def _find_parent_paragraph(self, root: ET.Element, element: ET.Element, namespaces: Dict) -> Optional[ET.Element]:
        """Find the parent paragraph element for a given element."""
        # Search all paragraphs and check if element is a descendant
        for para in root.findall('.//w:p', namespaces):
            # Check if the element is within this paragraph
            if self._is_descendant(para, element):
                return para
        return None
    
    def _is_descendant(self, ancestor: ET.Element, descendant: ET.Element) -> bool:
        """Check if descendant is a descendant of ancestor."""
        for elem in ancestor.iter():
            if elem == descendant:
                return True
        return False
    
    def _find_run_after_deletion(self, paragraph_elem: ET.Element, deletion_elem: ET.Element, doc: Document, namespaces: Dict):
        """Find a specific run that comes after a deletion element to attach comment to."""
        # Get all children of the paragraph
        children = list(paragraph_elem)
        
        try:
            # Find the index of the deletion element
            del_idx = children.index(deletion_elem)
            
            # Look for the next run element after the deletion
            for i in range(del_idx + 1, len(children)):
                child = children[i]
                if child.tag.endswith('}r'):  # It's a run element
                    # Get text from this run to find it in python-docx
                    run_text_elem = child.find('.//w:t', namespaces)
                    if run_text_elem is not None and run_text_elem.text:
                        # Find this run in python-docx by matching text
                        for para in doc.paragraphs:
                            for run in para.runs:
                                if run_text_elem.text[:20] in run.text or run.text[:20] in run_text_elem.text:
                                    return run
                    # If we found a run element but couldn't match it, use the first run after deletion
                    # Find paragraph in python-docx
                    para_text = ''.join([elem.text or '' for elem in paragraph_elem.findall('.//w:t', namespaces)])
                    for para in doc.paragraphs:
                        if para_text[:50] in para.text and para.runs:
                            # Return the first run that's not part of the deletion
                            return para.runs[0] if para.runs else None
        except (ValueError, IndexError):
            pass
        
        return None
    
    def _find_runs_for_redline_insertion(self, ins_elem: ET.Element, redline_text: str, doc: Document, root: ET.Element, namespaces: Dict):
        """Find ONLY the python-docx runs that are within the redline insertion element.
        
        CRITICAL: This method ONLY finds runs that are actually part of the <w:ins> element.
        It does NOT match regular document text outside the redline.
        """
        target_runs = []
        
        if not redline_text or not redline_text.strip():
            return target_runs
        
        # CRITICAL: Get the exact XML runs that are WITHIN the <w:ins> element
        # These are the only runs we want to match
        xml_runs_in_redline = ins_elem.findall('.//w:r', namespaces)
        
        if not xml_runs_in_redline:
            print(f"    ⚠ No XML runs found within redline element")
            return target_runs
        
        print(f"    Found {len(xml_runs_in_redline)} XML run(s) within redline element")
        
        # Get the parent paragraph of the redline element in XML
        # We need to find this paragraph in python-docx to get the actual run objects
        parent_para_xml = self._find_parent_paragraph(root, ins_elem, namespaces)
        
        if not parent_para_xml:
            print(f"    ⚠ Could not find parent paragraph in XML")
            return target_runs
        
        # Get all text from the parent paragraph to find it in python-docx
        para_text_elements = parent_para_xml.findall('.//w:t', namespaces)
        para_text = ''.join([elem.text or '' for elem in para_text_elements])
        
        # Find the matching paragraph in python-docx
        target_para = None
        for para in doc.paragraphs:
            # Match by text content
            if para_text[:100] in para.text or para.text[:100] in para_text:
                target_para = para
                break
        
        if not target_para or not target_para.runs:
            print(f"    ⚠ Could not find matching paragraph in python-docx")
            return target_runs
        
        # Now match each XML run from the redline to python-docx runs
        # We'll match by extracting text from each XML run and finding the corresponding python-docx run
        for xml_run in xml_runs_in_redline:
            # Get text from this XML run
            xml_text_elems = xml_run.findall('.//w:t', namespaces)
            xml_run_text = ''.join([elem.text or '' for elem in xml_text_elems if elem.text])
            
            if not xml_run_text.strip():
                continue
            
            # Find the matching python-docx run in the target paragraph
            # Match by text content - be precise
            for run in target_para.runs:
                # Only match if this run's text is part of the XML run text or vice versa
                # AND we haven't already added this run
                if (xml_run_text.strip() in run.text or run.text.strip() in xml_run_text) and run not in target_runs:
                    # Additional validation: make sure this run's text is actually in the redline
                    if run.text.strip() in redline_text or redline_text[:50] in run.text:
                        target_runs.append(run)
                        print(f"      Matched XML run '{xml_run_text[:30]}...' to python-docx run")
                        break
        
        # If we couldn't match by exact text, try a more lenient approach but still be strict
        if not target_runs:
            print(f"    ⚠ Could not match XML runs to python-docx runs by exact text")
            # Try matching by position - find runs in the paragraph that are likely the redline
            # Get the position of the redline element in the paragraph
            para_children = list(parent_para_xml)
            try:
                redline_idx = para_children.index(ins_elem)
                # The redline runs should be around this position
                # But we can't directly map XML positions to python-docx positions
                # So we'll use text matching with the redline text
                redline_words = redline_text.split()[:5]  # First 5 words
                if redline_words:
                    search_phrase = ' '.join(redline_words)
                    for run in target_para.runs:
                        if search_phrase in run.text and run not in target_runs:
                            target_runs.append(run)
                            if len(target_runs) >= 5:  # Limit to avoid selecting too much
                                break
            except (ValueError, IndexError):
                pass
        
        print(f"    Final: Found {len(target_runs)} python-docx run(s) matching redline")
        return target_runs
    
    def _get_run_properties_from_element(self, element: ET.Element, paragraph_elem: ET.Element, namespaces: Dict) -> Optional[ET.Element]:
        """Extract run properties (font, size, style) from an element or nearby runs.
        
        This ensures auto-redlines match the document's font and style.
        """
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # First, try to get rPr from runs within the target element
        for run in element.findall(f'.//{w_ns}r', namespaces):
            rpr = run.find(f'{w_ns}rPr')
            if rpr is not None:
                # Deep copy the rPr element
                import copy
                return copy.deepcopy(rpr)
        
        # If not found, look for runs in the paragraph
        for run in paragraph_elem.findall(f'.//{w_ns}r', namespaces):
            rpr = run.find(f'{w_ns}rPr')
            if rpr is not None:
                import copy
                return copy.deepcopy(rpr)
        
        return None
    
    def _insert_tracked_deletion(self, paragraph_elem: ET.Element, target_elem: ET.Element, text_to_delete: str, namespaces: Dict, author: str = 'RedLine Agent') -> None:
        """Insert a tracked deletion to strike out the counterparty's unacceptable text.
        
        Args:
            paragraph_elem: The parent paragraph XML element
            target_elem: The redline element to insert after (the counterparty's change)
            text_to_delete: The text being marked as deleted
            namespaces: XML namespaces
            author: Author name for the tracked change
        """
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # Create deletion element (tracked change)
        del_elem = ET.Element(f'{w_ns}del')
        del_elem.set(f'{w_ns}id', '0')
        del_elem.set(f'{w_ns}author', author)
        del_elem.set(f'{w_ns}date', datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
        
        # Create run element for the deleted text
        run_elem = ET.Element(f'{w_ns}r')
        
        # Copy run properties from document to match font/style
        rpr = self._get_run_properties_from_element(target_elem, paragraph_elem, namespaces)
        if rpr is not None:
            run_elem.append(rpr)
        
        # Create delText element (Word uses w:delText for deleted text content)
        del_text_elem = ET.Element(f'{w_ns}delText')
        del_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        del_text_elem.text = text_to_delete
        
        run_elem.append(del_text_elem)
        del_elem.append(run_elem)
        
        # Insert immediately after the target element within the paragraph
        children = list(paragraph_elem)
        try:
            target_idx = children.index(target_elem)
            paragraph_elem.insert(target_idx + 1, del_elem)
        except (ValueError, AttributeError):
            # Fallback: append if we can't locate the target
            paragraph_elem.append(del_elem)
    
    def _insert_auto_redline(self, paragraph_elem: ET.Element, target_elem: ET.Element, text: str, namespaces: Dict, is_restore: bool = False, author: str = 'RedLine Agent') -> None:
        """Insert an auto-redline (tracked change insertion) with the actual text.
        
        Matches the font and style of the surrounding document text.
        
        Args:
            paragraph_elem: The parent paragraph XML element
            target_elem: The redline element to insert after
            text: The actual text to insert (from playbook or restored original)
            namespaces: XML namespaces
            is_restore: If True, this is restoring deleted/replaced text
            author: Author name for the tracked change
        """
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # Create insertion element (tracked change)
        ins_elem = ET.Element(f'{w_ns}ins')
        ins_elem.set(f'{w_ns}id', '0')
        ins_elem.set(f'{w_ns}author', author)
        ins_elem.set(f'{w_ns}date', datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
        
        # Create run element for the inserted text
        run_elem = ET.Element(f'{w_ns}r')
        
        # Copy run properties from document to match font/style
        rpr = self._get_run_properties_from_element(target_elem, paragraph_elem, namespaces)
        if rpr is not None:
            run_elem.append(rpr)
        
        # Create text element with the actual text (no wrapper)
        text_elem = ET.Element(f'{w_ns}t')
        # Preserve spaces at boundaries
        text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        text_elem.text = text
        
        run_elem.append(text_elem)
        ins_elem.append(run_elem)
        
        # Insert immediately after the target element within the paragraph
        children = list(paragraph_elem)
        try:
            target_idx = children.index(target_elem)
            paragraph_elem.insert(target_idx + 1, ins_elem)
        except (ValueError, AttributeError):
            # Fallback: append if we can't locate the target
            paragraph_elem.append(ins_elem)
    
    def _insert_tracked_change_insertion(self, paragraph_elem: ET.Element, target_elem: ET.Element, text: str, namespaces: Dict, author: str = 'RedLine Agent') -> None:
        """Legacy method - redirects to _insert_auto_redline for backward compatibility."""
        self._insert_auto_redline(paragraph_elem, target_elem, text, namespaces, is_restore=False, author=author)
    
    def _insert_formatted_annotation(self, paragraph_elem: ET.Element, target_elem: ET.Element, annotation_text: str, risk_level: str, namespaces: Dict) -> None:
        """Insert formatted text annotation right after a redline element."""
        # Create a run for the annotation
        annotation_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        
        # Add run properties with formatting
        rpr = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        
        # Italic
        italic = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i')
        rpr.append(italic)
        
        # Risk-based color - Professional color scheme
        color_vals = {
            'High': 'C70039',  # Dark Red/Burgundy
            'Medium': 'FF8C00',  # Dark Orange/Amber
            'Low': '006400'  # Dark Green
        }
        color_val = color_vals.get(risk_level, '003366')  # Default Navy Blue
        
        color = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', color_val)
        rpr.append(color)
        
        # Get font size from surrounding content to match document font size
        font_size = self._get_font_size_from_element(target_elem, paragraph_elem, namespaces)
        sz = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
        sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', font_size)
        rpr.append(sz)
        
        annotation_run.append(rpr)
        
        # Add text with prefix
        text_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        text_elem.text = f" [AI Guidance - Risk: {risk_level}] {annotation_text}"
        annotation_run.append(text_elem)
        
        # Find position of target element and insert after it
        children = list(paragraph_elem)
        try:
            target_idx = children.index(target_elem)
            # Insert right after the target element
            paragraph_elem.insert(target_idx + 1, annotation_run)
        except (ValueError, AttributeError, IndexError):
            # Fallback: append to paragraph
            paragraph_elem.append(annotation_run)
    
    def _create_word_comment(self, paragraph_elem: ET.Element, target_elem: ET.Element, comments_root: ET.Element, comment_id: int, comment_text: str, risk_level: str, namespaces: Dict) -> None:
        """Create a native Word comment associated with a redline element.
        
        This creates:
        1. A comment entry in comments.xml
        2. Comment range markers in document.xml around the redline element
        """
        # CRITICAL: Ensure namespace is registered for 'w:' prefix
        ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
        
        # Create comment element in comments.xml
        comment_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
        comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
        comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'RedLine Agent')
        comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
        
        # CRITICAL: Create comment structure exactly as Word expects
        # Word requires: <w:comment><w:p><w:r><w:t>text</w:t></w:r></w:p></w:comment>
        
        # Ensure text is not None - Word requires actual text content
        if not comment_text or not comment_text.strip():
            comment_text = "Please review this change."
        
        # Clean and prepare text - replace multiple newlines with single newline
        clean_text = '\n'.join(line.strip() for line in comment_text.split('\n') if line.strip())
        if not clean_text:
            clean_text = "Please review this change."
        
        print(f"    Creating comment with text length: {len(clean_text)}", flush=True)
        print(f"    First 100 chars: {clean_text[:100]}", flush=True)
        
        # Create paragraph - REQUIRED
        comment_para = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
        
        # Add paragraph properties - REQUIRED by Word
        ppr = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        comment_para.append(ppr)
        
        # Split text into lines for multi-line comments
        text_lines = clean_text.split('\n')
        
        for line_idx, line in enumerate(text_lines):
            # Create a run for each line
            text_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            
            # Create text element - CRITICAL: This must have actual text content
            text_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            
            # Set the text - CRITICAL: text must be a string, not None
            # Also ensure special characters are properly handled (lxml will escape them)
            if line:
                # Ensure text is a string and not None
                text_elem.text = str(line) if line is not None else ' '
            else:
                # Empty line - use space
                text_elem.text = ' '
            
            # Add text element to run
            text_run.append(text_elem)
            
            # Add run to paragraph
            comment_para.append(text_run)
            
            # Add line break between lines (except after last line)
            if line_idx < len(text_lines) - 1:
                br = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                comment_para.append(br)
        
        # CRITICAL: Verify we have at least one run with text before adding to comment
        runs = comment_para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r', namespaces)
        has_text = False
        for run in runs:
            text_elems = run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', namespaces)
            for text_elem in text_elems:
                if text_elem.text and text_elem.text.strip():
                    has_text = True
                    break
            if has_text:
                break
        
        if not has_text:
            print(f"    ⚠ WARNING: No text found in comment paragraph! Creating fallback...", flush=True)
            # Fallback: create a simple run with text
            text_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            text_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text_elem.text = clean_text[:500] if clean_text else "Please review this change."
            text_run.append(text_elem)
            comment_para.append(text_run)
        
        # Add paragraph to comment element
        comment_elem.append(comment_para)
        
        # CRITICAL: Final verification - check that text is actually in the XML
        all_text_elems = comment_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', namespaces)
        total_text = ''
        for text_elem in all_text_elems:
            if text_elem.text:
                total_text += text_elem.text
        
        if total_text.strip():
            print(f"    ✓ Comment text verified in XML: '{total_text[:80]}...' (total length: {len(total_text)})", flush=True)
        else:
            print(f"    ✗ ERROR: Comment text is EMPTY in XML after creation!", flush=True)
            print(f"    Debug: Found {len(all_text_elems)} text elements", flush=True)
            for idx, te in enumerate(all_text_elems):
                print(f"      Text element {idx}: text={repr(te.text)}, tail={repr(te.tail)}", flush=True)
        
        comments_root.append(comment_elem)
        
        # Insert comment range markers in the document
        # Word comments need: commentRangeStart, commentReference (in a run), commentRangeEnd
        # The order and placement is critical
        # CRITICAL: For insertions, we want to highlight ONLY the redline content
        # For deletions, we attach to the paragraph containing the deletion
        
        # Find position of target element in paragraph
        children = list(paragraph_elem)
        try:
            target_idx = children.index(target_elem)
            
            # For insertions (<w:ins>), wrap the entire insertion element
            # CRITICAL: Insert markers in correct order and position
            # Word requires: commentRangeStart, [content], commentReference (in run), commentRangeEnd
            if target_elem.tag.endswith('}ins'):
                # Insertion: commentRangeStart before the <w:ins> element
                comment_range_start = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                comment_range_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                paragraph_elem.insert(target_idx, comment_range_start)
                
                # After inserting commentRangeStart, target_elem is now at target_idx + 1
                # Find the position right after the <w:ins> element ends
                # The <w:ins> element is now at target_idx + 1, so we insert after it
                ins_end_idx = target_idx + 1
                
                # Create commentReference in a run - this must come AFTER the insertion content
                # CRITICAL: The run containing commentReference MUST have text or a space, otherwise Word won't display it
                comment_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                comment_ref = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                comment_ref.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                comment_run.append(comment_ref)
                # Add a space text element so Word recognizes the run
                space_text = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                space_text.text = ' '
                comment_run.append(space_text)
                
                # Insert commentReference run right after the <w:ins> element
                # We need to find where the ins element actually ends in the children list
                # Since we inserted commentRangeStart, the ins is now at target_idx + 1
                # Find the next sibling after the ins element
                current_children = list(paragraph_elem)
                ins_elem_idx = None
                for i, child in enumerate(current_children):
                    if child == target_elem:
                        ins_elem_idx = i
                        break
                
                if ins_elem_idx is not None:
                    # Insert commentReference run after the ins element
                    paragraph_elem.insert(ins_elem_idx + 1, comment_run)
                    # Insert commentRangeEnd after the commentReference
                    comment_range_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                    comment_range_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                    paragraph_elem.insert(ins_elem_idx + 2, comment_range_end)
                else:
                    # Fallback: append at end
                    paragraph_elem.append(comment_run)
                    comment_range_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                    comment_range_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                    paragraph_elem.append(comment_range_end)
                
            elif target_elem.tag.endswith('}del'):
                # Deletion: place markers around the deletion element
                comment_range_start = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                comment_range_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                paragraph_elem.insert(target_idx, comment_range_start)
                
                # Find the next run after deletion to attach comment to
                comment_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                comment_ref = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                comment_ref.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                comment_run.append(comment_ref)
                # Add a space text element so Word recognizes the run
                space_text = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                space_text.text = ' '
                comment_run.append(space_text)
                # Insert after the <w:del> element
                paragraph_elem.insert(target_idx + 2, comment_run)
                
                comment_range_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                comment_range_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                paragraph_elem.insert(target_idx + 3, comment_range_end)
            else:
                # Unknown element type - use fallback
                raise ValueError(f"Unknown element type: {target_elem.tag}")
            
        except (ValueError, AttributeError, IndexError) as e:
            print(f"    ⚠ Could not find exact position for comment markers: {e}")
            # Fallback: append at end of paragraph
            comment_range_start = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
            comment_range_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
            paragraph_elem.append(comment_range_start)
            
            comment_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            comment_ref = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
            comment_ref.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
            comment_run.append(comment_ref)
            # Add a space text element so Word recognizes the run
            space_text = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            space_text.text = ' '
            comment_run.append(space_text)
            paragraph_elem.append(comment_run)
            
            comment_range_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
            comment_range_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
            paragraph_elem.append(comment_range_end)
    
    def _add_comment_annotation(self, paragraph, comment_text: str, risk_level: str) -> None:
        """Add a comment annotation to a paragraph."""
        from docx.shared import RGBColor, Pt
        
        # Get font size from paragraph to match document font size
        paragraph_font_size = None
        if paragraph.runs:
            # Get size from first run that has a font size
            for run in paragraph.runs:
                if run.font.size:
                    paragraph_font_size = run.font.size
                    break
        
        # Default to paragraph style size if no run size found
        if paragraph_font_size is None and paragraph.style and paragraph.style.font.size:
            paragraph_font_size = paragraph.style.font.size
        
        # Default to 12pt if nothing found
        if paragraph_font_size is None:
            paragraph_font_size = Pt(12)
        
        # Add a line break before comment
        paragraph.add_run().add_break()
        
        # Add comment with formatting based on risk level - Professional color scheme
        risk_colors = {
            'High': RGBColor(199, 0, 57),  # Dark Red/Burgundy
            'Medium': RGBColor(255, 140, 0),  # Dark Orange/Amber
            'Low': RGBColor(0, 100, 0)  # Dark Green
        }
        
        color = risk_colors.get(risk_level, RGBColor(0, 51, 102))  # Default Navy Blue
        
        # Add comment prefix - match paragraph font size
        prefix_run = paragraph.add_run(f"[AI Analysis - Risk: {risk_level}] ")
        prefix_run.bold = True
        prefix_run.font.color.rgb = color
        prefix_run.font.size = paragraph_font_size
        
        # Add comment text - match paragraph font size
        comment_run = paragraph.add_run(comment_text)
        comment_run.italic = True
        comment_run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray
        comment_run.font.size = paragraph_font_size
    
    def insert_comments_google(self, analyses: List[Dict]) -> None:
        """Insert comments into Google Doc."""
        if not self.service or not self.doc_id:
            raise ValueError("Google Docs service not initialized")
        
        # Get document to find text positions
        doc = self.service.documents().get(documentId=self.doc_id).execute()
        
        requests = []
        
        for analysis in analyses:
            redline = analysis['redline']
            comment_text = analysis.get('comment_text', analysis.get('response', ''))
            
            # Find the position of the redline text in the document
            text_to_find = redline.get('text', '')[:100]  # First 100 chars
            
            # Search for text position
            position = self._find_text_position(doc, text_to_find)
            
            if position is not None:
                # Create comment request
                request = {
                    'createComment': {
                        'location': {
                            'index': position
                        },
                        'comment': {
                            'content': [
                                {
                                    'text': comment_text
                                }
                            ]
                        }
                    }
                }
                requests.append(request)
        
        # Batch execute requests
        if requests:
            self.service.documents().batchUpdate(
                documentId=self.doc_id,
                body={'requests': requests}
            ).execute()
    
    def _find_text_position(self, doc: Dict, search_text: str) -> Optional[int]:
        """Find the position of text in Google Doc."""
        if 'body' not in doc or 'content' not in doc['body']:
            return None
        
        current_index = 1  # Google Docs uses 1-based indexing
        
        for element in doc['body']['content']:
            if 'paragraph' in element:
                para = element['paragraph']
                if 'elements' in para:
                    for elem in para['elements']:
                        if 'textRun' in elem:
                            text = elem['textRun'].get('content', '')
                            if search_text in text:
                                # Find the exact position within the text
                                text_pos = text.find(search_text)
                                return current_index + text_pos
                            current_index += len(text)
        
        return None
    
    def create_summary_document(self, analyses: List[Dict], output_path: str) -> None:
        """Create a separate summary document with all analyses."""
        doc = Document()
        doc.add_heading('Redline Analysis Summary', 0)
        
        for idx, analysis in enumerate(analyses, 1):
            redline = analysis['redline']
            
            doc.add_heading(f'Redline #{idx}', level=1)
            doc.add_paragraph(f"Type: {redline['type']}")
            doc.add_paragraph(f"Text: {redline.get('text', 'N/A')}")
            doc.add_paragraph(f"Author: {redline.get('author', 'Unknown')}")
            doc.add_paragraph(f"Date: {redline.get('date', 'Unknown')}")
            
            doc.add_paragraph(f"Assessment: {analysis.get('assessment', 'N/A')}")
            doc.add_paragraph(f"Risk Level: {analysis.get('risk_level', 'N/A')}")
            doc.add_paragraph(f"Recommended Action: {analysis.get('response', 'N/A')}")
            if analysis.get('fallbacks'):
                doc.add_paragraph(f"Fallback/Alternative: {analysis.get('fallbacks', 'N/A')}")
            
            doc.add_paragraph('')  # Blank line
        
        doc.save(output_path)

